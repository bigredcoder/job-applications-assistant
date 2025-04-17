#!/usr/bin/env python3
"""
Create a sample Word document for testing
"""

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_resume():
    # Create a new document
    doc = docx.Document()
    
    # Add a title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Brian Robison")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Add contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("San Francisco, CA | brian.robison@example.com | (555) 123-4567 | linkedin.com/in/brianrobison")
    
    # Add summary
    doc.add_heading("Summary", level=2)
    summary = doc.add_paragraph()
    summary.add_run("Experienced software engineer with 8+ years of experience in full-stack development, specializing in Python, JavaScript, and cloud technologies. Proven track record of delivering scalable, high-performance applications and leading engineering teams.")
    
    # Add skills
    doc.add_heading("Skills", level=2)
    skills = doc.add_paragraph()
    skills.add_run("Languages: ").bold = True
    skills.add_run("Python, JavaScript, TypeScript, SQL, HTML/CSS\n")
    skills.add_run("Frameworks: ").bold = True
    skills.add_run("React, Django, Flask, Express.js, Next.js\n")
    skills.add_run("Cloud & DevOps: ").bold = True
    skills.add_run("AWS (EC2, S3, Lambda), Docker, Kubernetes, CI/CD pipelines\n")
    skills.add_run("Databases: ").bold = True
    skills.add_run("PostgreSQL, MongoDB, Redis\n")
    skills.add_run("Tools: ").bold = True
    skills.add_run("Git, JIRA, Confluence, Slack, VS Code")
    
    # Add experience
    doc.add_heading("Experience", level=2)
    
    # Job 1
    job1_title = doc.add_paragraph()
    job1_title.add_run("Senior Software Engineer | TechCorp Inc. | Jan 2020 - Present").bold = True
    
    job1_bullets = doc.add_paragraph(style='List Bullet')
    job1_bullets.add_run("Led development of a microservices architecture that improved system reliability by 35%")
    
    job1_bullets = doc.add_paragraph(style='List Bullet')
    job1_bullets.add_run("Architected and implemented RESTful APIs serving 1M+ daily requests")
    
    job1_bullets = doc.add_paragraph(style='List Bullet')
    job1_bullets.add_run("Mentored junior developers and conducted code reviews to maintain code quality")
    
    job1_bullets = doc.add_paragraph(style='List Bullet')
    job1_bullets.add_run("Reduced infrastructure costs by 25% through cloud optimization strategies")
    
    job1_bullets = doc.add_paragraph(style='List Bullet')
    job1_bullets.add_run("Implemented automated testing that increased code coverage from 65% to 92%")
    
    # Job 2
    job2_title = doc.add_paragraph()
    job2_title.add_run("Software Engineer | DataSystems LLC | Mar 2017 - Dec 2019").bold = True
    
    job2_bullets = doc.add_paragraph(style='List Bullet')
    job2_bullets.add_run("Developed full-stack web applications using React, Node.js, and PostgreSQL")
    
    job2_bullets = doc.add_paragraph(style='List Bullet')
    job2_bullets.add_run("Created data visualization dashboards that helped clients increase operational efficiency")
    
    # Save the document
    doc.save("baseline.docx")
    print("Created sample resume: baseline.docx")

if __name__ == "__main__":
    create_sample_resume()
