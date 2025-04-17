#!/usr/bin/env python3
"""
Utility functions for handling different file formats
"""

import os
import io
from pathlib import Path
from typing import Union, BinaryIO, Optional

def extract_text_from_file(file_path_or_obj: Union[str, Path, BinaryIO, bytes]) -> str:
    """
    Extract text from various file formats including .txt, .md, .doc, and .docx
    
    Args:
        file_path_or_obj: Path to the file or file-like object
        
    Returns:
        Extracted text content as string
    """
    # Determine file extension
    if isinstance(file_path_or_obj, (str, Path)):
        file_path = str(file_path_or_obj)
        file_ext = os.path.splitext(file_path)[1].lower()
    elif hasattr(file_path_or_obj, 'name'):
        # For file-like objects with a name attribute (like StreamingBody)
        file_ext = os.path.splitext(file_path_or_obj.name)[1].lower()
    else:
        # Default to txt if we can't determine the extension
        file_ext = '.txt'
    
    # Handle different file formats
    if file_ext in ['.txt', '.md']:
        # Plain text or markdown
        if isinstance(file_path_or_obj, (str, Path)):
            with open(file_path_or_obj, 'r', encoding='utf-8') as f:
                return f.read()
        elif hasattr(file_path_or_obj, 'read'):
            # File-like object
            if hasattr(file_path_or_obj, 'seek'):
                file_path_or_obj.seek(0)
            
            if isinstance(file_path_or_obj, io.TextIOBase):
                return file_path_or_obj.read()
            else:
                return file_path_or_obj.read().decode('utf-8')
        elif isinstance(file_path_or_obj, bytes):
            return file_path_or_obj.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file object type: {type(file_path_or_obj)}")
    
    elif file_ext == '.docx':
        # Word .docx format
        try:
            import docx
            
            if isinstance(file_path_or_obj, (str, Path)):
                doc = docx.Document(file_path_or_obj)
            else:
                # For file-like objects or bytes
                if isinstance(file_path_or_obj, bytes):
                    file_obj = io.BytesIO(file_path_or_obj)
                else:
                    # Ensure we're at the beginning of the file
                    if hasattr(file_path_or_obj, 'seek'):
                        file_path_or_obj.seek(0)
                    file_obj = file_path_or_obj
                
                doc = docx.Document(file_obj)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            return '\n'.join(full_text)
        
        except ImportError:
            raise ImportError("python-docx library is required to read .docx files. Install it with 'pip install python-docx'")
    
    elif file_ext == '.doc':
        # Word .doc format (older)
        try:
            import textract
            
            if isinstance(file_path_or_obj, (str, Path)):
                return textract.process(file_path_or_obj).decode('utf-8')
            else:
                # For file-like objects or bytes, we need to save to a temporary file
                import tempfile
                
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    if isinstance(file_path_or_obj, bytes):
                        temp_file.write(file_path_or_obj)
                    else:
                        # Ensure we're at the beginning of the file
                        if hasattr(file_path_or_obj, 'seek'):
                            file_path_or_obj.seek(0)
                        temp_file.write(file_path_or_obj.read())
                    
                    temp_path = temp_file.name
                
                try:
                    return textract.process(temp_path).decode('utf-8')
                finally:
                    # Clean up the temporary file
                    os.unlink(temp_path)
        
        except ImportError:
            raise ImportError("textract library is required to read .doc files. Install it with 'pip install textract'")
    
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")
