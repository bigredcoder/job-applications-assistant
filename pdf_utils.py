#!/usr/bin/env python3
"""
Utility functions for PDF generation
"""

import base64
import io
import os
import tempfile
from typing import Optional, Tuple

# Try to import PDF generation libraries
try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    reportlab_available = True
except ImportError:
    reportlab_available = False

# Try to import docx generation libraries
try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    docx_available = True
except ImportError:
    docx_available = False

def markdown_to_html(markdown_text: str) -> str:
    """
    Convert markdown to HTML with styling.

    Args:
        markdown_text: Markdown text to convert

    Returns:
        HTML string with styling
    """
    try:
        # Try to use markdown if available
        import markdown
        html = markdown.markdown(markdown_text, extensions=['extra'])
    except ImportError:
        # Fallback to a simple conversion if markdown is not available
        html = markdown_text.replace('\n\n', '</p><p>')
        html = f"<p>{html}</p>"
        html = html.replace('# ', '<h1>').replace('\n## ', '</h1><h2>')
        html = html.replace('## ', '<h2>').replace('\n### ', '</h2><h3>')
        html = html.replace('### ', '<h3>').replace('\n#### ', '</h3><h4>')
        html = html.replace('#### ', '<h4>').replace('\n##### ', '</h4><h5>')
        html = html.replace('**', '<strong>').replace('**', '</strong>')
        html = html.replace('*', '<em>').replace('*', '</em>')

    # Add minimal CSS for better formatting
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                margin: 1in;
                font-size: 11pt;
            }}
            h1, h2, h3, h4 {{
                color: #333;
                margin-top: 1.2em;
                margin-bottom: 0.5em;
            }}
            h1 {{ font-size: 18pt; }}
            h2 {{ font-size: 16pt; }}
            h3 {{ font-size: 14pt; }}
            h4 {{ font-size: 12pt; }}
            p {{ margin: 0.5em 0; }}
            ul, ol {{ margin: 0.5em 0; padding-left: 1.5em; }}
            a {{ color: #0066cc; text-decoration: none; }}
            hr {{ border: 1px solid #ddd; }}
            .header {{ text-align: center; margin-bottom: 1em; }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """

    return styled_html

def get_markdown_download_link(markdown_text: str, filename: str) -> str:
    """
    Generate a download link for a markdown file.

    Args:
        markdown_text: Markdown text to download
        filename: Name of the file to download

    Returns:
        HTML string with download link
    """
    # Encode the markdown text as base64
    markdown_base64 = base64.b64encode(markdown_text.encode('utf-8')).decode('utf-8')

    # Create a download link
    href = f'<a href="data:text/markdown;base64,{markdown_base64}" download="{filename}">Download {filename}</a>'

    return href

def get_html_download_link(markdown_text: str, filename: str) -> str:
    """
    Generate a download link for an HTML file created from markdown.

    Args:
        markdown_text: Markdown text to convert to HTML
        filename: Name of the file to download

    Returns:
        HTML string with download link
    """
    # Convert markdown to HTML
    html = markdown_to_html(markdown_text)

    # Encode the HTML as base64
    html_base64 = base64.b64encode(html.encode('utf-8')).decode('utf-8')

    # Create a download link
    href = f'<a href="data:text/html;base64,{html_base64}" download="{filename.replace(".pdf", ".html")}">Download as HTML</a>'

    return href

def markdown_to_docx_base64(markdown_text: str) -> Optional[str]:
    """
    Convert markdown to DOCX and return as base64 encoded string.

    Args:
        markdown_text: Markdown text to convert

    Returns:
        Base64 encoded DOCX or None if conversion failed
    """
    if not docx_available:
        print("python-docx is not available. Cannot generate DOCX.")
        return None

    try:
        # Create a new document
        document = docx.Document()

        # Set margins (1 inch on all sides)
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Process markdown into paragraphs
        lines = markdown_text.split('\n')
        in_list = False
        list_items = []

        for line in lines:
            # Handle headings
            if line.startswith('# '):
                heading = document.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = document.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = document.add_heading(line[4:], level=3)
            # Handle list items
            elif line.startswith('- ') or line.startswith('* '):
                if not in_list:
                    in_list = True
                list_items.append(line[2:])
            # Handle empty lines
            elif line.strip() == "":
                if in_list and list_items:
                    # Add the list items
                    for item in list_items:
                        p = document.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(item)
                    list_items = []
                    in_list = False
                document.add_paragraph()
            # Handle normal paragraphs
            else:
                if in_list:
                    # Add any pending list items
                    for item in list_items:
                        p = document.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(item)
                    list_items = []
                    in_list = False

                # Handle bold and italic text
                p = document.add_paragraph()
                text = line

                # Simple handling of bold and italic
                # This is a simplified approach - a more robust parser would be better
                while '**' in text or '*' in text:
                    if '**' in text:
                        parts = text.split('**', 2)
                        p.add_run(parts[0])
                        bold_run = p.add_run(parts[1])
                        bold_run.bold = True
                        text = parts[2] if len(parts) > 2 else ''
                    elif '*' in text:
                        parts = text.split('*', 2)
                        p.add_run(parts[0])
                        italic_run = p.add_run(parts[1])
                        italic_run.italic = True
                        text = parts[2] if len(parts) > 2 else ''

                # Add any remaining text
                if text:
                    p.add_run(text)

        # Add any remaining list items
        if in_list and list_items:
            for item in list_items:
                p = document.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(item)

        # Save the document to a bytes buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Get the docx data
        docx_data = buffer.getvalue()
        buffer.close()

        # Return the base64 encoded docx
        return base64.b64encode(docx_data).decode("utf-8")

    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return None

def get_docx_download_link(markdown_text: str, filename: str) -> Optional[str]:
    """
    Generate a download link for a DOCX file created from markdown.

    Args:
        markdown_text: Markdown text to convert to DOCX
        filename: Name of the file to download

    Returns:
        HTML string with download link or None if conversion failed
    """
    docx_base64 = markdown_to_docx_base64(markdown_text)
    if docx_base64:
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_base64}" download="{filename.replace(".pdf", ".docx")}">Download as Word (.docx)</a>'
        return href
    return None

def get_pdf_download_link(markdown_text: str, filename: str) -> Optional[str]:
    """
    Generate a download link for a PDF created from markdown.

    Args:
        markdown_text: Markdown text to convert to PDF
        filename: Name of the file to download

    Returns:
        HTML string with download link or None if conversion failed
    """
    pdf_base64 = markdown_to_pdf_base64(markdown_text)
    if pdf_base64:
        href = f'<a href="data:application/pdf;base64,{pdf_base64}" download="{filename}">Download as PDF</a>'
        return href
    return None

def markdown_to_pdf_base64(markdown_text: str) -> Optional[str]:
    """
    Convert markdown to PDF using ReportLab and return as base64 encoded string.

    Args:
        markdown_text: Markdown text to convert

    Returns:
        Base64 encoded PDF or None if conversion failed
    """
    if not reportlab_available:
        print("ReportLab is not available. Cannot generate PDF.")
        return None

    try:
        # Create a buffer for the PDF
        buffer = io.BytesIO()

        # Create the PDF document
        doc = SimpleDocTemplate(buffer, pagesize=LETTER,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)

        # Get styles
        styles = getSampleStyleSheet()

        # Create custom styles
        heading1_style = ParagraphStyle(
            name='CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12
        )

        heading2_style = ParagraphStyle(
            name='CustomHeading2',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=10
        )

        heading3_style = ParagraphStyle(
            name='CustomHeading3',
            parent=styles['Heading3'],
            fontSize=14,
            spaceAfter=8
        )

        normal_style = ParagraphStyle(
            name='CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6
        )

        # Process markdown into paragraphs
        story = []

        # Simple markdown parsing
        lines = markdown_text.split('\n')
        current_paragraph = ""

        for line in lines:
            if line.startswith('# '):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, normal_style))
                    current_paragraph = ""
                story.append(Paragraph(line[2:], heading1_style))
            elif line.startswith('## '):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, normal_style))
                    current_paragraph = ""
                story.append(Paragraph(line[3:], heading2_style))
            elif line.startswith('### '):
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, normal_style))
                    current_paragraph = ""
                story.append(Paragraph(line[4:], heading3_style))
            elif line.strip() == "":
                if current_paragraph:
                    story.append(Paragraph(current_paragraph, normal_style))
                    current_paragraph = ""
                story.append(Spacer(1, 0.2 * inch))
            else:
                # Handle bold and italic - more carefully
                # First, escape any XML/HTML special characters
                line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                # Handle bold text with proper pairing
                bold_parts = line.split('**')
                if len(bold_parts) > 1:
                    new_line = bold_parts[0]
                    for i in range(1, len(bold_parts)):
                        if i % 2 == 1:  # Opening bold
                            new_line += '<b>'
                        else:  # Closing bold
                            new_line += '</b>'
                        new_line += bold_parts[i]
                    line = new_line

                # Handle italic text with proper pairing
                italic_parts = line.split('*')
                if len(italic_parts) > 1:
                    new_line = italic_parts[0]
                    for i in range(1, len(italic_parts)):
                        if i % 2 == 1:  # Opening italic
                            new_line += '<i>'
                        else:  # Closing italic
                            new_line += '</i>'
                        new_line += italic_parts[i]
                    line = new_line

                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line

        # Add the last paragraph if there is one
        if current_paragraph:
            story.append(Paragraph(current_paragraph, normal_style))

        # Build the PDF
        doc.build(story)

        # Get the PDF data
        pdf_data = buffer.getvalue()
        buffer.close()

        # Return the base64 encoded PDF
        return base64.b64encode(pdf_data).decode("utf-8")

    except Exception as e:
        print(f"Error generating PDF with ReportLab: {e}")
        return None
